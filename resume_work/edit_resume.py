from pathlib import Path

from docx import Document


SOURCE = Path(r"D:\my-coding\TaskForge\resume_work\source.docx")
OUTPUT = Path(r"C:\Users\christina\Desktop\AI_Agent_中文简历_TaskForge版.docx")


def set_single(paragraph, text: str) -> None:
    """Replace a one-run paragraph while retaining its formatting."""
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""


def set_bullet(paragraph, text: str) -> None:
    """Replace the body of an existing bullet and retain the blue marker."""
    if len(paragraph.runs) < 2:
        set_single(paragraph, "•  " + text)
        return
    paragraph.runs[0].text = "•  "
    paragraph.runs[1].text = text
    for run in paragraph.runs[2:]:
        run.text = ""


def set_title(paragraph, title: str, date: str) -> None:
    """Replace a project title while preserving title/date styles."""
    if len(paragraph.runs) < 2:
        set_single(paragraph, f"{title}   {date}")
        return
    paragraph.runs[0].text = title
    paragraph.runs[1].text = f"   {date}"
    for run in paragraph.runs[2:]:
        run.text = ""


def set_labeled(paragraph, label: str, body: str) -> None:
    """Replace a label/body paragraph while retaining the bold label."""
    if len(paragraph.runs) < 2:
        set_single(paragraph, label + body)
        return
    paragraph.runs[0].text = label
    paragraph.runs[1].text = body
    for run in paragraph.runs[2:]:
        run.text = ""


doc = Document(SOURCE)

# Keep personal fields editable; improve the role positioning and skills.
set_single(doc.paragraphs[1], "求职方向：AI Agent / LLM 应用工程师")
set_bullet(doc.paragraphs[6], "Agent Runtime：LangGraph 有界状态机 + 受限 ReAct；Tool Gateway、宿主授权执行、checkpoint/幂等；固定 DAG + handoff")
set_bullet(doc.paragraphs[7], "RAG 与长文档：BM25 / BGE 向量 / RRF / MiniLM rerank；QASPER、TAT-QA、MultiHop 评测；证据级 grounding 与无证据拒答")
set_bullet(doc.paragraphs[8], "工程与安全：Python/FastAPI/Pydantic/Vue3/SQLite/Qdrant；工具白名单、风险分级、人工接管、CI、可复现评测与回归门禁")

# TaskForge: align the headline with the actual repository and current evidence.
set_title(doc.paragraphs[10], "TaskForge｜权限受控、可恢复的通用 Agent Runtime", "2026.07—至今")
set_labeled(
    doc.paragraphs[11],
    "背景 / 栈：",
    "模型仅提交结构化 ToolRequest，宿主负责权限、幂等、checkpoint 与证据验证；固定多角色审查 DAG。Python·FastAPI·Pydantic·Vue3·SQLite·Qdrant·DeepSeek·CI",
)
set_bullet(
    doc.paragraphs[12],
    "构建 Provider-neutral AgentRuntime + Tool Gateway：JSON Schema 白名单、风险分级、审批、幂等、断点续跑；464 项自动化测试通过，ruff 与前端构建纳入 CI",
)
set_bullet(
    doc.paragraphs[13],
    "实现企业审查多角色 DAG：检索回执签发一次性 receipt，claim 按 verified/proposed 分级，跨角色 handoff 仅携带已验证事实；无证据时阻断模型结论",
)
set_bullet(
    doc.paragraphs[14],
    "建立四场景检索评测与回归门禁：Recall@10 分别为 QASPER 73.4%（独立验证 72.2%）、TAT-QA 99.0%、MultiHop 92.0%、PDF 100%；阻断局部调优导致的跨路由退化",
)

# PatchPilot: make the safety/control story more direct.
set_labeled(
    doc.paragraphs[16],
    "背景 / 栈：",
    "限制模型自由读写，宿主控制路径、补丁与测试；LangGraph 修复闭环。Python·LangGraph·Typer·GitPython·Pytest",
)
set_bullet(
    doc.paragraphs[17],
    "LangGraph 状态图编排修复闭环（规划→探索→diff→安全预检→补丁→回归→归因→重试≤3 轮）；工作树全量 249 passed / 2 skipped",
)
set_bullet(
    doc.paragraphs[18],
    "受控探索：模型仅 5 个只读工具，宿主持有路径/补丁/测试；确定性安全门（3 文件/120 行、禁删、pytest 白名单+超时）",
)
set_bullet(
    doc.paragraphs[19],
    "失败归因（Provider 故障 vs 补丁质量）驱动重试，持久化 trajectory/checkpoint；对接 QuixBugs、SWE-bench Lite（300 case）",
)

# TripBound: retain the project but sharpen its product value.
set_labeled(
    doc.paragraphs[21],
    "背景 / 栈：",
    "外部数据波动与 POI 幻觉；以事实边界与可回退 Provider 保障可控行程。FastAPI·Vue3·SQLAlchemy·Redis·AMap/Geoapify",
)
set_bullet(
    doc.paragraphs[22],
    "Guarded ReAct 事实边界：模型仅引用候选 ID、Python 白名单物化行程，杜绝幻觉 POI；Provider 降级链（AMap MCP + Geoapify），不补造票价/班次",
)
set_bullet(
    doc.paragraphs[23],
    "旅行客服 RAG：词法召回 + 可插拔 Embedding/RRF + 引用验证/无证据拒答；支持多 Agent 并行检索",
)
set_bullet(
    doc.paragraphs[24],
    "588 项测试通过；POI 门禁评测覆盖 HitRate@K、MRR、nDCG，并包含负例控制，避免指标虚高",
)

doc.save(OUTPUT)
print(OUTPUT)
